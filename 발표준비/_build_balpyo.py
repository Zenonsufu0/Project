# -*- coding: utf-8 -*-
"""기말 발표자료 생성기 → B04_기말발표자료.docx

강의자료 #11(기말발표) 발표자료 구성 요구사항 반영:
표지(+동료평가 평점) / 팀원별 담당 파트 / 2차 확장 내용 /
기획서·설계 수정 사유·경위 / 검사 중 주요 오류 검출 / 보강사항 / (선택)2차 확장계획
A4 세로, 가독성 우선(심미성 무관).
"""
from pathlib import Path

import docx
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "발표준비" / "B04_기말발표자료.docx"

KOR = "맑은 고딕"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
OWNER = RGBColor(0x70, 0x30, 0xA0)
GOOD = RGBColor(0x1E, 0x6B, 0x2F)
BAD = RGBColor(0xC0, 0x00, 0x00)


def set_kor_font(run, size=None, bold=None, color=None):
    run.font.name = KOR
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), KOR)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def para(doc, *segments, align=None, space_after=6, size=11):
    """segments: str 또는 (text, dict) ; dict에 bold/color/size."""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    for seg in segments:
        if isinstance(seg, tuple):
            text, opt = seg
        else:
            text, opt = seg, {}
        r = p.add_run(text)
        set_kor_font(r, size=opt.get("size", size), bold=opt.get("bold"), color=opt.get("color"))
    return p


def heading(doc, text, size=15, color=ACCENT, before=14, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_kor_font(r, size=size, bold=True, color=color)
    # 밑줄 경계선
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "1F4E79")
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def owner_tag(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_kor_font(r, size=9.5, bold=True, color=OWNER)
    return p


def bullets(doc, items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("· ")
        set_kor_font(r, size=size, bold=True, color=ACCENT)
        if isinstance(it, tuple):
            head, rest = it
            r1 = p.add_run(head); set_kor_font(r1, size=size, bold=True)
            r2 = p.add_run(rest); set_kor_font(r2, size=size)
        else:
            r1 = p.add_run(it); set_kor_font(r1, size=size)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    # borders
    tblPr = t._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), "999999")
        borders.append(el)
    tblPr.append(borders)
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = c.paragraphs[0].add_run(h)
        set_kor_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "1F4E79")
        c._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = cells[j].paragraphs[0].add_run(str(val))
            set_kor_font(run, size=9.5)
            if ri % 2 == 1:
                shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "EAF1F8")
                cells[j]._tc.get_or_add_tcPr().append(shd)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    return t


def build():
    doc = docx.Document()
    # A4 세로
    sec = doc.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(1.8)
    sec.top_margin = sec.bottom_margin = Cm(1.8)
    # Normal 폰트
    normal = doc.styles["Normal"]
    normal.font.name = KOR; normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), KOR)

    # ===== 표지 =====
    for _ in range(2):
        doc.add_paragraph()
    para(doc, ("기말 발표자료", {"size": 26, "bold": True, "color": ACCENT}), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, ("건국 수강신청 시뮬레이터 — KU Course Registration Simulator", {"size": 13, "bold": True}),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, ("Python CLI · 전공기초프로젝트 (2026-1)", {"size": 11, "color": RGBColor(0x55, 0x55, 0x55)}),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    para(doc, ("팀명: B04 (5인)", {"size": 13, "bold": True}), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    heading(doc, "팀원 명단 및 동료평가 평점 (잠정 합의)", size=13, before=6)
    table(doc,
          ["학번", "이름", "주 담당 모듈", "동료평가 평점"],
          [["202312363", "이강준", "main.py (UI·흐름)", "3"],
           ["202312349", "박성수", "auth.py (인증)", "3"],
           ["202111376", "조서현", "storage.py (파일·무결성)", "3"],
           ["202312364", "이도현", "student_service.py (수강신청)", "3"],
           ["202312356", "신경환", "admin_service.py (관리자)", "3"]],
          widths=[3.2, 2.4, 6.5, 3.2])
    para(doc, ("⚠ 동료평가 평점은 발표 전 팀원 합의로 확정하여 위 표를 수정할 것 (표지 필수 항목). "
               "평점 {3,2,1,0} 중 하나, 자기 자신 포함.", {"size": 9.5, "color": BAD}), space_after=2)

    doc.add_page_break()

    # ===== 1. 팀원별 담당 파트 =====
    heading(doc, "1. 팀원별 담당 파트")
    para(doc, ("모듈 단위로 기획·구현·검사를 분담했다. 기획은 전원 협업, 구현·검사는 각 모듈 담당자가 수행하고 전원이 교차 검토했다.",
               {"size": 10.5}))
    table(doc,
          ["모듈 / 파트", "주요 책임", "구현", "검사(TC)"],
          [["main.py", "진입점·CLI 흐름·메뉴·무결성 호출", "이강준", "이강준"],
           ["auth.py", "로그인·회원가입·ID/PW/이름/학년 검증", "박성수", "박성수"],
           ["storage.py", "10개 파일 입출력·무결성 검사", "조서현", "조서현"],
           ["student_service.py", "수강신청 11단계·취소·시간표·학점", "이도현", "이도현"],
           ["admin_service.py", "학생·강의·강의실 관리·기간 설정", "신경환", "신경환"],
           ["models.py (공통)", "Student/Course/Schedule/Classroom 등 데이터 클래스", "전원 협업", "—"],
           ["기획 / 2차 기획서·설계문서", "요구사항·데이터 구조 설계", "전원 협업", "전원 교차검토"]],
          widths=[4.3, 6.6, 2.3, 2.3])
    para(doc, ("총 검사 TC: 1차 129개 중 18개 삭제(2차 수행불가/결과변경) → 선별 승계 111개 + 2차 신규·변경 53개 = 164개.", {"size": 10, "bold": True}))
    owner_tag(doc, "{발표: 이강준}")

    # ===== 2. 2차 확장 내용 =====
    heading(doc, "2. 2차 확장 내용 (추가·변경된 기능)")
    para(doc, ("교수 질의응답을 반영해 채택한 0~5번 기능을 설계문서(원판) 기준으로 구현했다. 핵심은 ‘강의 정보 정규화(스케줄·강의실 분리)’와 "
               "‘수강신청 제약 강화(학년·학과·학기·선수과목)’이다.", {"size": 10.5}))
    bullets(doc, [
        ("강의실 신설 — ", "classrooms.txt 추가, 관리자 ‘강의실 관리(10번)’ 메뉴(목록·등록). 강의 정원은 배정 강의실 중 최소 좌석 수를 초과할 수 없음."),
        ("스케줄 분리 — ", "요일·시각·강의실을 courses.txt에서 schedules.txt로 분리. 한 강의가 복수 스케줄(요일별 다른 강의실) 가능. 시간표 충돌 판정을 스케줄 기반으로 변경."),
        ("선수과목 — ", "prerequisites.txt 추가. 수강신청 시 기이수(completed) 기준으로 선수과목 이수 여부 검증, 미이수 시 과목명까지 안내."),
        ("학년·학과·학기 제한 — ", "students.txt에 학년, courses.txt에 학기·제한학년·제한학과 추가. 수강신청 검사를 7단계 → 11단계로 확장(학기·학년·학과·선수과목)."),
        ("강의 수정 항목별 재설계 — ", "통합 수정을 항목별로 분리. 수강신청 기간 시작 후에도 정원·교강사·강의실은 변경 가능(제약 검증), 과목명·학점·요시는 변경 불가."),
    ])
    para(doc, ("[1차 미진 → 마저 구현] 데이터 파일의 행/필드 앞뒤 공백을 문법 오류로 감지하는 처리, 무결성 검사 메시지 형식 정합을 2차에서 마무리.",
               {"size": 10, "color": RGBColor(0x33, 0x55, 0x99)}))
    owner_tag(doc, "{기획: 전원 / 발표: 신경환·이도현}")

    # ===== 3. 기획서·설계문서 수정 사유·경위 =====
    heading(doc, "3. 기획서·설계문서 수정 사유·경위")
    para(doc, ("2차 기획서 ‘설계 기준판’ 제출 (원판 → 설계 기준판). 설계 단계에서 발견한 명세 누락을 보완했다.", {"size": 10.5}))
    table(doc,
          ["수정 문서 / 항목", "수정 사유", "발견 경위"],
          [["2차 기획서 설계기준판\n6.14.2 강의 수정",
            "과목명·학점 변경의 ‘수강신청 기간 중 가능 여부’가 명시되지 않아 구현 판정 기준이 정해지지 않는 누락 → 두 항목을 ‘기간 시작 후 불가’로 명시",
            "설계 단계에서 항목별 수정 규칙을 코드로 옮기던 중, 정원·교강사·강의실과 달리 과목명·학점의 기간 제한이 빠진 것을 발견"],
           ["설계문서 (원판)\n9.2 항목별 제한 규칙",
            "상태 변경은 강의 삭제·활성화와 일관되게 기간 무관 허용으로 정리(원판 유지). 정원 증감·강의실 변경의 좌석 제약을 명문화",
            "구현·검사 중 정원/강의실 경계 조건(좌석 수, 현재 신청 인원)에서 판정 기준 필요성 확인"]],
          widths=[4.5, 7.0, 4.0])
    para(doc, ("※ 최종 수정판(2차 기획서 최종수정판/설계문서 수정판)은 추가 수정이 필요한 경우에만 제출(해당 시 본 절에 갱신).",
               {"size": 9.5, "color": RGBColor(0x55, 0x55, 0x55)}))
    owner_tag(doc, "{기획: 전원 / 발표: 신경환}")

    # ===== 4. 검사 중 주요 오류 검출 =====
    heading(doc, "4. 검사 중 주요 오류 검출 (TC와 함께)")
    para(doc, ("2차 구현·검사 과정에서 테스트 케이스로 검출·교정한 주요 사항. 각 항목은 ‘발견하게 만든 TC’와 함께 기재한다.", {"size": 10.5}))
    table(doc,
          ["TC", "검출 내용 (검사로 발견한 문제 / 경계)", "교정·확인 결과"],
          [["6.12-5",
            "스케줄이 2개인 과목의 학점이 2회 합산되어 총 신청 학점이 부풀려지는 버그",
            "current_credits를 과목(분반) 단위로 중복 제거 → 총 3학점(6 아님) 확인"],
           ["6.9-20 / 6.9-21",
            "충돌 판정을 스케줄 기반으로 옮기면서, [시작,종료) 반열림 경계(종료==시작)에서 오탐 우려",
            "겹침 조건 start<end’ and start’<end 로 통일 → 경계 비충돌, 실제 겹침만 차단 확인"],
           ["6.9-18",
            "선수과목 미이수 시 어떤 과목이 부족한지 안내 부재",
            "미이수 선수과목명 목록 출력 ‘[1002] 자료구조’ 확인"],
           ["6.14-17 / 6.14-22",
            "정원이 배정 강의실 좌석 수를 초과해도 등록/증가가 통과될 위험",
            "최소 좌석 강의실 기준 상한 검증 → ‘가장 작은 강의실(…석) 초과 불가’ 차단 확인"],
           ["6.14-23",
            "정원을 현재 신청 인원보다 작게 줄이면 데이터 모순 발생",
            "현재 enrolled 인원 미만 감소 차단 → ‘현재 신청 인원(N명) 이상’ 확인"],
           ["5.5-15 / 5.5-16",
            "스키마 변경(students 7필드·config 3필드) 후 구(舊) 형식 파일이 조용히 통과될 위험",
            "필드 수 검사로 ‘필드 수 오류’ 검출 → 무결성 단계에서 차단 확인"]],
          widths=[2.4, 8.1, 5.0])
    owner_tag(doc, "{검사·발표: 조서현(5.5)·이도현(6.9/6.12)·신경환(6.14)}")

    # ===== 5. 보강사항 =====
    heading(doc, "5. 보강사항 (마감 후 발견 · ‘아차’ 사항)")
    para(doc, ("제출 마감 후 점검에서 발견한 사항들. 발표자료에 명시하여 감점 폭을 줄이고, 강사가 못 찾은 문제는 선제 보고한다.", {"size": 10.5}))
    bullets(doc, [
        ("[검사보고서] 1차 6.14 TC 일부 대체 — ", "강의 수정이 통합→항목별로 재설계되어 1차 검사보고서의 일부 6.14 TC는 그대로 적용되지 않는다. 2차 검사보고서 17절의 6.14-16~29로 재검증하여 대체했다."),
        ("[기획서] 문구 vs 규칙 불일치 — ", "기획서 6.14.2의 오류 문구가 ‘수강신청 기간 중’으로 되어 있으나 실제 규칙은 ‘기간 시작 후(중+종료 후)’ 차단이다. 구현은 ‘시작 이후’ 의미로 통일했고, 문구는 최종수정판에서 ‘기간 시작 이후’로 정정 예정."),
        ("[기본 데이터] 수강신청 기간 폭 — ", "config 기본값이 시작=종료=오늘(기획서 5.2.5 명세)이라 기본 상태로는 수강신청 가능 구간이 매우 짧다. 시연 시 관리자 9번으로 기간을 넓혀 시연한다."),
        ("[기획서] 충돌 용어 불일치 — ", "기획서 용어 절은 '시간표 충돌'을 '스케줄 충돌'로 대체한다고 선언했으나 6.9절 6단계 오류 문구는 '시간표 충돌'로 남아 있다. 구현·설계문서·검사보고서는 '스케줄 충돌'로 통일했으며, 최종수정판 제출 시 문구 정정 대상이다."),
        ("[설계 정합] 강의 수정 진입 — ", "설계문서 7.10대로 강의 수정 ‘진입’은 기간과 무관하게 허용하고 항목별로만 제한하도록 구현했다(1차의 진입 자체 차단과 다름). 일관성 확인 완료."),
    ])
    owner_tag(doc, "{발표: 이도현}")

    # ===== 6. (선택) 2차 확장 계획 / 차기 =====
    heading(doc, "6. (선택) 차기 확장 아이디어")
    para(doc, ("교수 피드백에서 제외 권고된 항목(수강바구니·재수강 평점·졸업 시뮬레이션·이수구분)은 훈련 적합도/작업량을 고려해 도입하지 않았다. "
               "추가 시간이 허락되면 ‘강의실 시간대 중복 배정 방지’(같은 강의실 동시간 2강의 금지)를 우선 검토한다.", {"size": 10.5}))
    owner_tag(doc, "{발표: 전원}")

    # ===== 부록: 발표 진행 메모 =====
    heading(doc, "부록. 발표 진행 메모", size=12, before=14)
    bullets(doc, [
        "실시간 카카오톡 그룹콜 · 팀당 20분(학생 발표 15분 이내 + 질의응답). 강사가 그룹콜을 걸 때까지 대기.",
        "발언 시 항상 자기 이름 먼저. 분업한 부분은 담당자가 직접 발표·답변.",
        "발표 자료·문서·코드를 각자 화면에 띄우고 진행(화면 공유 없음).",
        "공지 시간보다 최대 20분 지연 가능 — 대기.",
    ], size=10)

    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    out = build()
    print("saved:", out)
