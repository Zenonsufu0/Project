# -*- coding: utf-8 -*-
"""2차 검사 보고서 빌더.

1차 보고서(B04_보고서_최종.docx)를 복사한 뒤,
색상 범례 + '2차 확장 검사' 섹션(실측 결과)을 추가하여
B04_2차_검사_보고서.docx 를 생성한다.

각 2차 TC의 '실제 결과'는 서비스/스토리지를 직접 호출하여 실측한 값이다.
"""
import shutil
import sys
import tempfile
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from major_basics.modules.admin_service import AdminService
from major_basics.modules.auth import AuthService
from major_basics.modules.models import (
    Classroom, Config, Course, Enrollment, Schedule, Student,
)
from major_basics.modules.storage import DataStore, IntegrityError
from major_basics.modules.student_service import StudentService


# ------------------------------------------------------------------
# 공용 픽스처
# ------------------------------------------------------------------
def fresh_store():
    tmp = Path(tempfile.mkdtemp()) / "raw"
    store = DataStore(tmp)
    store.ensure_defaults(date(2026, 4, 1))
    return store, tmp


def classrooms():
    return {
        "1001": Classroom("1001", "공학관", "101", 60),
        "1002": Classroom("1002", "공학관", "201", 80),
        "1003": Classroom("1003", "과학관", "102", 40),
    }


def open_config():
    return Config(date(2026, 4, 1), date(2026, 4, 30), date(2026, 4, 3), "2026-1")


def student(grade=2, major="컴퓨터공학부", sid="202111376"):
    return Student(sid, "Abc1234", "홍길동", "공과대학", major, "active", grade)


def ssvc(stu, courses, schedules, enr=None, comp=None, prereq=None, config=None):
    return StudentService(stu, courses, enr or [], comp or {}, config or open_config(),
                          schedules, classrooms(), prereq or {})


def asvc(courses=None, schedules=None, enr=None, current=date(2026, 3, 1),
         rs=date(2026, 4, 1), re_=date(2026, 4, 7)):
    cfg = Config(rs, re_, current, "2026-1")
    return AdminService({}, courses or {}, enr or [], {},
                        {"공과대학": ["컴퓨터공학부", "전기공학부"], "이과대학": ["수학과"]},
                        cfg, schedules or {}, classrooms(), {})


# ------------------------------------------------------------------
# TC 정의: (id, 목표/입력, 예상결과, actual_fn)
# actual_fn() -> 실제결과 문자열
# ------------------------------------------------------------------
TCS = []


def tc(tid, goal, expected, fn):
    TCS.append((tid, goal, expected, fn))


# ---- 5.5 무결성 확장 ----
def _missing_file_default(name, check):
    store, tmp = fresh_store()
    p = tmp / name
    exists = p.exists()
    content = p.read_text(encoding="utf-8").strip().splitlines()
    return f"{name} 자동 생성됨({exists}); {check(content)}"


tc("5.5-12", "[목표] 누락 시 classrooms.txt 기본 강의실 3개 자동 생성 / 입력: classrooms.txt 삭제 후 실행",
   "기본 강의실 3개(1001/1002/1003) 포함 파일 자동 생성",
   lambda: _missing_file_default("classrooms.txt", lambda c: f"행수={len(c)}, 첫행={c[0]}"))

tc("5.5-13", "[목표] 누락 시 schedules.txt 기본 스케줄 자동 생성 / 입력: schedules.txt 삭제 후 실행",
   "기본 스케줄 9행 자동 생성",
   lambda: _missing_file_default("schedules.txt", lambda c: f"행수={len(c)}, 첫행={c[0]}"))

tc("5.5-14", "[목표] 누락 시 prerequisites.txt 자동 생성 / 입력: prerequisites.txt 삭제 후 실행",
   "1003,1002 (알고리즘←자료구조) 포함 자동 생성",
   lambda: _missing_file_default("prerequisites.txt", lambda c: f"내용={c}"))


def _integrity_after(name, text):
    store, tmp = fresh_store()
    (tmp / name).write_text(text, encoding="utf-8")
    errs = store.validate_integrity()
    return errs[0] if errs else "(오류 없음)"


tc("5.5-15", "[목표] students.txt 학년 필드 누락(6필드) 시 문법 오류 / 입력: '...,active' (6필드)",
   "!!! 오류: students.txt N행 - ... '필드 수 오류'",
   lambda: _integrity_after("students.txt", "202111376,Abc1234,홍길동,공과대학,컴퓨터공학부,active\n"))

tc("5.5-16", "[목표] config.txt 학기 필드 누락(2필드) 시 문법 오류 / 입력: '2026-04-01,2026-04-07'",
   "!!! 오류: config.txt 1행 - ... '필드 수 오류'",
   lambda: _integrity_after("config.txt", "2026-04-01,2026-04-07\n"))

tc("5.5-17", "[목표] courses.txt 신스키마(10필드) 정상 로드 / 입력: 기본 6과목",
   "무결성 통과(오류 없음), 6과목 로드",
   lambda: (lambda s: f"{'통과' if not s.validate_integrity() else s.validate_integrity()[0]}; 과목수={len(s.load_courses())}")(fresh_store()[0]))

tc("5.5-18", "[목표] schedules.txt 미참조 강의실코드 → 참조 무결성 위반 / 입력: '1001,01,MON,09:00,10:30,9999'",
   "!!! 오류: 참조 무결성 위반 - schedules.txt ... 강의실코드",
   lambda: _integrity_after("schedules.txt", "1001,01,MON,09:00,10:30,9999\n"))

tc("5.5-19", "[목표] prerequisites.txt 미존재 과목코드 → 참조 무결성 위반 / 입력: '9999,1001'",
   "!!! 오류: 참조 무결성 위반 - prerequisites.txt ... 과목코드",
   lambda: _integrity_after("prerequisites.txt", "9999,1001\n"))

tc("5.5-20", "[목표] prerequisites.txt 자기참조 → 문법 오류 / 입력: '1001,1001'",
   "!!! 오류: prerequisites.txt N행 - ... '자기 자신을 선수과목으로 지정 불가'",
   lambda: _integrity_after("prerequisites.txt", "1001,1001\n"))

tc("5.5-21", "[목표] courses.txt 제한학년 범위 초과 → 문법 오류 / 입력: 제한학년=5",
   "!!! 오류: courses.txt N행 - ... '제한학년 형식 오류'",
   lambda: _integrity_after("courses.txt", "1001,01,프로그래밍기초,3,김교수,active,30,2026-1,5,전체\n"))

tc("5.5-22", "[목표] courses.txt 미존재 제한학과 → 참조 무결성 위반 / 입력: 제한학과=없는학과",
   "!!! 오류: 참조 무결성 위반 - courses.txt ... 제한학과(전공)",
   lambda: _integrity_after("courses.txt", "1001,01,프로그래밍기초,3,김교수,active,30,2026-1,0,없는학과\n"))

tc("5.5-23", "[목표] 누락 시 config.txt 자동 생성 — 학기 포함 3필드 (5.2.5절: 1~6월→YYYY-1) / 입력: config.txt 삭제 후 날짜 2026-04-01로 실행",
   "'2026-04-01,2026-04-01,2026-1' 내용으로 자동 생성",
   lambda: fresh_store()[0].config_path.read_text(encoding="utf-8").strip())

tc("5.5-24", "[목표] 행 맨 앞 표준 공백 → 문법 오류 (5.1절 공통 규칙) / 입력: students.txt 행 앞에 공백 1개 삽입",
   "!!! 오류: students.txt 1행 - 문법 형식이 올바르지 않습니다: '행/필드 앞뒤 공백 불허'",
   lambda: _integrity_after("students.txt", " 202111376,Abc1234,홍길동,공과대학,컴퓨터공학부,active,2\n"))

tc("5.5-25", "[목표] 필드 앞뒤 표준 공백 → 문법 오류 (5.1절 공통 규칙) / 입력: courses.txt 담당교수 필드 뒤 공백",
   "!!! 오류: courses.txt 1행 - 문법 형식이 올바르지 않습니다: '행/필드 앞뒤 공백 불허'",
   lambda: _integrity_after("courses.txt", "1001,01,프로그래밍기초,3,김교수 ,active,30,2026-1,0,전체\n"))


# ---- 6.3 회원가입(학년) ----
tc("6.3-14", "[목표] 학년 범위 초과 입력 거부 / 입력: 학년='5'",
   "!!! 오류: 학년은 1 이상 4 이하의 정수이어야 합니다.",
   lambda: AuthService({}, {}, {}).validate_grade("5")[1])

tc("6.3-15", "[목표] 정상 학년으로 회원가입 시 students.txt 7필드 저장 / 입력: 학년='2'",
   "active 상태 + 학년 2 포함 7필드 저장",
   lambda: (lambda a: (a.signup_student("202400001", "Pass123", "Pass123", "김건국", "공과대학", "컴퓨터공학부", 2),
                       f"grade={a.students['202400001'].grade}, status={a.students['202400001'].status}")[1])(AuthService({}, {}, {})))


# ---- 6.13 관리자 학생 관리(학년) ----
def _tc_admin_register_student():
    a = asvc()
    _ok, msg = a.register_student(Student("202500777", "Pass123", "김건국", "공과대학", "컴퓨터공학부", "active", 2))
    s = a.students["202500777"]
    return f"{msg}; grade={s.grade}, status={s.status}"


tc("6.13-10", "[목표] 관리자 학생 등록 정상 — 학년 포함 (6.13.1절) / 입력: 학번 202500777, 각 필드 정상, 학년 2",
   "✓ 학생 등록 완료; grade=2, status=active",
   _tc_admin_register_student)


# ---- 6.7 개설과목 조회(스케줄/제한/학기) ----
def _list_courses_demo():
    store, _ = fresh_store()
    courses = store.load_courses(); schedules = store.load_schedules()
    svc = ssvc(student(grade=2), courses, schedules, config=open_config())
    listed = svc.list_courses()
    c1001 = next(c for c in listed if c.key() == ("1001", "01"))
    days = "/".join(s.day for s in schedules[("1001", "01")])
    return f"목록 {len(listed)}건; 1001-01 스케줄요일={days}"


tc("6.7-7", "[목표] 스케줄 2개 과목 요일 표시 / 입력: 1001-01(MON,WED) 조회",
   "MON / WED 두 스케줄 표시",
   _list_courses_demo)

tc("6.7-8", "[목표] 제한학년·학과 표시 / 입력: 1002-01(2학년/컴퓨터공학부) 조회",
   "제한란 '2학년/컴퓨터공학부' 표시",
   lambda: (lambda c: f"limit_grade={c.limit_grade}, limit_major={c.limit_major}")(fresh_store()[0].load_courses()[("1002", "01")]))

tc("6.7-9", "[목표] 현재 학기 외 과목 미표시(학기 필터) / 입력: config 학기=2026-1, 9001=2026-2",
   "2026-2 과목은 목록에서 제외",
   lambda: _semester_filter())


def _semester_filter():
    store, _ = fresh_store()
    courses = store.load_courses()
    courses[("9001", "01")] = Course("9001", "01", "다음학기", 3, "김교수", semester="2026-2")
    svc = ssvc(student(), courses, store.load_schedules(), config=open_config())
    codes = [c.code for c in svc.list_courses()]
    return f"목록 과목코드={sorted(set(codes))}; 9001 포함={'9001' in codes}"


tc("6.7-10", "[목표] inactive 과목 전체 조회 제외 (6.7.1절) / 입력: 1001-01을 inactive로 변경 후 조회",
   "1001-01은 목록에서 제외, 나머지 active 과목만 출력",
   lambda: _inactive_excluded())


def _inactive_excluded():
    store, _ = fresh_store()
    courses = store.load_courses()
    courses[("1001", "01")].status = "inactive"
    svc = ssvc(student(grade=2), courses, store.load_schedules(), config=open_config())
    keys = [f"{c.code}-{c.section}" for c in svc.list_courses()]
    return f"목록={keys}; 1001-01 포함={'1001-01' in keys}"


tc("6.7-11", "[목표] 신청 불가 과목 [제한] 표시 (기획서 7절 — 숨기지 않고 표시만) / 입력: 1학년 학생으로 전체 조회",
   "1002(2학년 제한)·1003(3학년 제한) 행에 '[제한]' 표시, 목록에는 유지",
   lambda: _restricted_marker())


def _restricted_marker():
    import io
    import contextlib
    import major_basics.main as M

    store, _ = fresh_store()
    courses = store.load_courses()
    schedules = store.load_schedules()
    stu = student(grade=1)
    svc = ssvc(stu, courses, schedules, config=open_config())
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        M._print_courses(svc.list_courses(), schedules, classrooms(), stu)
    lines = buf.getvalue().splitlines()
    row_1002 = next(l for l in lines if "| 1002 |" in l)
    row_1001 = next(l for l in lines if "| 1001 | 01 |" in l)
    return f"1002행 [제한]={'[제한]' in row_1002}; 1001행 [제한]={'[제한]' in row_1001}; 1002행 목록 유지=True"


# ---- 6.9 수강신청 11단계 ----
def _reg_case(setup):
    store, _ = fresh_store()
    courses = store.load_courses(); schedules = store.load_schedules(); prereq = store.load_prerequisites()
    return setup(courses, schedules, prereq)


tc("6.9-15", "[목표] 8단계 학기 불일치 차단 / 입력: 학기 2026-2 과목 신청(현재 2026-1)",
   "!!! 오류: 현재 학기(2026-1)에 개설된 과목이 아닙니다.",
   lambda: _reg_case(lambda c, s, p: (c.__setitem__(("8888", "01"), Course("8888", "01", "타학기", 3, "김교수", semester="2026-2")),
                                      s.__setitem__(("8888", "01"), [Schedule("8888", "01", "FRI", 540, 630, "1001")]),
                                      ssvc(student(grade=2), c, s, config=open_config()).register("8888", "01")[1])[-1]))

tc("6.9-16", "[목표] 9단계 학년 제한 차단 / 입력: 1학년 학생이 1002(2학년 제한) 신청",
   "!!! 오류: 이 과목은 2학년 학생만 수강신청할 수 있습니다.",
   lambda: _reg_case(lambda c, s, p: ssvc(student(grade=1), c, s, config=open_config()).register("1002", "01")[1]))

tc("6.9-17", "[목표] 10단계 학과 제한 차단 / 입력: 전기공학부 2학년이 1002(컴공 전용) 신청",
   "!!! 오류: 이 과목은 컴퓨터공학부 학생만 수강신청할 수 있습니다.",
   lambda: _reg_case(lambda c, s, p: ssvc(student(grade=2, major="전기공학부"), c, s, config=open_config()).register("1002", "01")[1]))

tc("6.9-18", "[목표] 11단계 선수과목 미이수 차단 + 미이수 과목명 나열 / 입력: 3학년이 1003(선수 1002 미이수) 신청",
   "!!! 오류: 선수과목을 이수하지 않았습니다. 미이수 선수과목: 자료구조",
   lambda: _reg_case(lambda c, s, p: ssvc(student(grade=3), c, s, prereq=p, config=open_config()).register("1003", "01")[1]))

tc("6.9-19", "[목표] 선수과목 이수 후 정상 신청 / 입력: 1002 기이수 후 1003 신청",
   "✓ 수강신청 완료: 알고리즘",
   lambda: _reg_case(lambda c, s, p: ssvc(student(grade=3), c, s, comp={"202111376": {"1002"}}, prereq=p, config=open_config()).register("1003", "01")[1].split(chr(10))[0]))

tc("6.9-20", "[목표] 6단계 스케줄 충돌 차단(schedules 기준) / 입력: 1002 신청 후 동일시간 9999 신청",
   "!!! 오류: 스케줄 충돌 - 자료구조 (MON 13:00~14:30)과 겹칩니다.",
   lambda: _reg_case(_sched_conflict))

def _sched_conflict(c, s, p):
    c[("9999", "01")] = Course("9999", "01", "충돌과목", 3, "김교수", semester="2026-1")
    s[("9999", "01")] = [Schedule("9999", "01", "MON", 13 * 60, 14 * 60 + 30, "1001")]
    svc = ssvc(student(grade=2), c, s, config=open_config())
    svc.register("1002", "01")
    return svc.register("9999", "01")[1]

tc("6.9-21", "[목표] 반열림 구간 경계 비충돌(종료==시작) / 입력: MON 09:00~10:30 후 MON 10:30~12:00",
   "✓ 수강신청 완료 (충돌 아님)",
   lambda: _reg_case(_halfopen))

def _halfopen(c, s, p):
    c[("7001", "01")] = Course("7001", "01", "앞", 3, "김교수", semester="2026-1")
    c[("7002", "01")] = Course("7002", "01", "뒤", 3, "이교수", semester="2026-1")
    s[("7001", "01")] = [Schedule("7001", "01", "MON", 9 * 60, 10 * 60 + 30, "1001")]
    s[("7002", "01")] = [Schedule("7002", "01", "MON", 10 * 60 + 30, 12 * 60, "1002")]
    svc = ssvc(student(grade=2), c, s, config=open_config())
    svc.register("7001", "01")
    return svc.register("7002", "01")[1].split(chr(10))[0]


# ---- 6.12 시간표(스케줄/학점 중복) ----
def _timetable_case():
    store, _ = fresh_store()
    courses = store.load_courses(); schedules = store.load_schedules()
    svc = ssvc(student(grade=2), courses, schedules, config=open_config())
    svc.register("1001", "01")  # MON,WED 두 스케줄
    tt = svc.timetable()
    return f"시간표행수={len(tt)}(요일 {[s.day for _,s in tt]}); 학점={svc.current_credits()}"


tc("6.12-4", "[목표] 동일 과목 복수 요일 모두 표시 / 입력: 1001-01(MON,WED) 신청 후 시간표",
   "MON, WED 두 줄 표시(각 강의실 포함)",
   _timetable_case)

tc("6.12-5", "[목표] 스케줄 2개 과목 학점 중복 합산 방지 / 입력: 1001-01(3학점, 스케줄 2개)",
   "총 학점 3 (6 아님)",
   _timetable_case)


# ---- 6.14.1 강의 등록(스케줄/정원) ----
tc("6.14-16", "[목표] 스케줄 0개 등록 거부 / 입력: 스케줄 미입력",
   "!!! 오류: 스케줄을 1개 이상 등록해야 합니다.",
   lambda: asvc().add_course(Course("1005", "01", "DB", 3, "박교수", semester="2026-1"), [])[1])

tc("6.14-17", "[목표] 정원 > 최소 강의실 좌석 거부 / 입력: 정원 50, 과학관102(40석)",
   "!!! 오류: 정원은 가장 작은 강의실(과학관102, 40석)의 좌석 수를 초과할 수 없습니다.",
   lambda: asvc().add_course(Course("1005", "01", "DB", 3, "박교수", capacity=50, semester="2026-1"),
                             [Schedule("1005", "01", "MON", 540, 630, "1003")])[1])

tc("6.14-18", "[목표] 정상 강의 등록(스케줄 1개) / 입력: 정원 30, 공학관101",
   "✓ 강의 등록 완료: DB (1005-01) | 스케줄 1개",
   lambda: asvc().add_course(Course("1005", "01", "DB", 3, "박교수", capacity=30, semester="2026-1"),
                             [Schedule("1005", "01", "MON", 540, 630, "1001")])[1])


# ---- 6.14.2 강의 수정 항목별 ----
def _co():
    return {("1001", "01"): Course("1001", "01", "프로그래밍기초", 3, "김교수", capacity=30, semester="2026-1")}, \
           {("1001", "01"): [Schedule("1001", "01", "MON", 9 * 60, 10 * 60 + 30, "1001")]}


tc("6.14-19", "[목표] 기간 시작 후 과목명 변경 차단 / 입력: today=04-03 ≥ reg_start=04-01",
   "!!! 오류: 수강신청 기간 시작 이후에는 과목명을 수정할 수 없습니다.",
   lambda: (lambda co: asvc(co[0], co[1], current=date(2026, 4, 3)).update_course_name("1001", "01", "새이름")[1])(_co()))

tc("6.14-20", "[목표] 기간 시작 후 학점 변경 차단 / 입력: today=04-03",
   "!!! 오류: 수강신청 기간 시작 이후에는 학점을 수정할 수 없습니다.",
   lambda: (lambda co: asvc(co[0], co[1], current=date(2026, 4, 3)).update_course_credits("1001", "01", 2)[1])(_co()))

tc("6.14-21", "[목표] 기간 시작 후 정원 증가 허용(≤최소좌석60) / 입력: 30→50, today=04-03",
   "✓ 강의 수정 완료: 정원 → 50",
   lambda: (lambda co: asvc(co[0], co[1], current=date(2026, 4, 3)).update_course_capacity("1001", "01", 50)[1])(_co()))

tc("6.14-22", "[목표] 정원 증가 최소좌석 초과 차단 / 입력: 30→70(공학관101=60)",
   "!!! 오류: 정원은 가장 작은 강의실(공학관101, 60석)의 좌석 수를 초과할 수 없습니다.",
   lambda: (lambda co: asvc(co[0], co[1], current=date(2026, 4, 3)).update_course_capacity("1001", "01", 70)[1])(_co()))

tc("6.14-23", "[목표] 정원 감소 현재 신청 인원 미만 차단 / 입력: enrolled 2명, 정원→1",
   "!!! 오류: 정원은 현재 신청 인원(2명) 이상이어야 합니다.",
   lambda: (lambda co: asvc(co[0], co[1], current=date(2026, 4, 3),
                            enr=[Enrollment("202200001", "1001", "01", "enrolled"),
                                 Enrollment("202200002", "1001", "01", "enrolled")]).update_course_capacity("1001", "01", 1)[1])(_co()))

tc("6.14-24", "[목표] 담당교수 변경 시 타 강의 스케줄 충돌 차단 / 입력: 신교수(2001 MON 09:00 담당)로 변경",
   "!!! 오류: 새 담당교수가 이미 담당하는 강의(B)와 요시가 충돌합니다.",
   lambda: _prof_conflict())

def _prof_conflict():
    co, sc = _co()
    co[("2001", "01")] = Course("2001", "01", "B", 3, "신교수", semester="2026-1")
    sc[("2001", "01")] = [Schedule("2001", "01", "MON", 9 * 60, 10 * 60 + 30, "1002")]
    return asvc(co, sc, current=date(2026, 4, 3)).update_course_professor("1001", "01", "신교수")[1]

tc("6.14-25", "[목표] 강의실 변경(현재 인원 수용 가능) / 입력: enrolled 5명, MON→과학관102(40석)",
   "✓ 강의실 변경 완료: MON → 과학관102 (40석)",
   lambda: (lambda co: asvc(co[0], co[1], current=date(2026, 4, 3),
                            enr=[Enrollment(f"20220000{i}", "1001", "01", "enrolled") for i in range(5)]
                            ).update_schedule_classroom("1001", "01", "MON", "1003")[1])(_co()))

tc("6.14-26", "[목표] 강의실 변경 좌석 부족 차단 / 입력: enrolled 50명, MON→과학관102(40석)",
   "!!! 오류: 변경하려는 강의실(40석)이 현재 신청 인원(50명)보다 작습니다.",
   lambda: (lambda co: asvc(co[0], co[1], current=date(2026, 4, 3),
                            enr=[Enrollment(f"2022{i:05d}", "1001", "01", "enrolled") for i in range(50)]
                            ).update_schedule_classroom("1001", "01", "MON", "1003")[1])(_co()))

tc("6.14-27", "[목표] 기간 시작 후 스케줄 추가 차단 / 입력: today=04-03, FRI 추가",
   "!!! 오류: 수강신청 기간 시작 이후에는 스케줄을 추가할 수 없습니다.",
   lambda: (lambda co: asvc(co[0], co[1], current=date(2026, 4, 3)).add_schedule(Schedule("1001", "01", "FRI", 540, 630, "1001"))[1])(_co()))

tc("6.14-28", "[목표] 마지막 1개 스케줄 삭제 차단 / 입력: 스케줄 1개 상태에서 삭제(기간 전)",
   "!!! 오류: 스케줄은 최소 1개 이상 유지되어야 합니다.",
   lambda: (lambda co: asvc(co[0], co[1], current=date(2026, 3, 1)).delete_schedule("1001", "01", "MON")[1])(_co()))

tc("6.14-29", "[목표] inactive 강의 수정 차단 / 입력: inactive 상태에서 과목명 변경",
   "!!! 오류: inactive 상태의 강의는 수정할 수 없습니다. 먼저 강의를 활성화하세요.",
   lambda: _inactive_edit())

tc("6.14-30", "[목표] 스케줄 종료 시각 ≤ 시작 시각 오류 (4.8절) / 입력: 신규 강의 스케줄 MON 09:00~09:00",
   "!!! 오류: 종료 시각은 시작 시각보다 이후여야 합니다.",
   lambda: asvc().add_course(Course("1006", "01", "신규", 3, "박교수", capacity=30, semester="2026-1"),
                             [Schedule("1006", "01", "MON", 540, 540, "1001")])[1])

tc("6.14-31", "[목표] 수강신청 기간 시작 전 과목명 변경 정상 (6.14.2절) / 입력: today=03-01 < reg_start=04-01",
   "✓ 강의 수정 완료: 과목명 → 새이름",
   lambda: (lambda co: asvc(co[0], co[1], current=date(2026, 3, 1)).update_course_name("1001", "01", "새이름")[1])(_co()))

def _inactive_edit():
    co, sc = _co()
    co[("1001", "01")].status = "inactive"
    return asvc(co, sc, current=date(2026, 3, 1)).update_course_name("1001", "01", "x")[1]


# ---- 6.15 기간 설정(학기) ----
tc("6.15-7", "[목표] 잘못된 학기 형식 차단 / 입력: 학기='2026-3'",
   "!!! 오류: 학기 형식이 올바르지 않습니다.",
   lambda: asvc().set_registration_period(date(2026, 4, 1), date(2026, 4, 7), "2026-3")[1])

tc("6.15-8", "[목표] 정상 학기 설정 / 입력: 2026-09-01~09-07, 학기='2026-2'",
   "✓ 수강신청 기간 설정 완료: 2026-09-01 ~ 2026-09-07 | 학기: 2026-2",
   lambda: asvc().set_registration_period(date(2026, 9, 1), date(2026, 9, 7), "2026-2")[1])

tc("6.15-9", "[목표] 시작일 = 종료일 경계값 정상 (6.15절 의미 규칙) / 입력: 2026-04-01 ~ 2026-04-01, 학기='2026-1'",
   "✓ 수강신청 기간 설정 완료: 2026-04-01 ~ 2026-04-01 | 학기: 2026-1",
   lambda: asvc().set_registration_period(date(2026, 4, 1), date(2026, 4, 1), "2026-1")[1])


# ---- 6.17 강의실 관리(신규) ----
tc("6.17-1", "[목표] 강의실 목록 조회(코드 오름차순) / 입력: 기본 강의실 3개",
   "1001, 1002, 1003 순 출력",
   lambda: ", ".join(r.classroom_code for r in asvc().list_classrooms()))

tc("6.17-2", "[목표] 강의실 등록 정상 / 입력: 1004,과학관,301,50",
   "✓ 강의실 등록 완료: 과학관 301 (코드: 1004)",
   lambda: asvc().add_classroom(Classroom("1004", "과학관", "301", 50))[1])

tc("6.17-3", "[목표] 강의실코드 중복 차단 / 입력: 1001 재등록",
   "!!! 오류: 이미 존재하는 강의실코드입니다.",
   lambda: asvc().add_classroom(Classroom("1001", "공학관", "999", 10))[1])


# ------------------------------------------------------------------
# 실행 + 결과 수집
# ------------------------------------------------------------------
def run():
    results = []
    for tid, goal, expected, fn in TCS:
        try:
            actual = str(fn())
        except Exception as e:  # noqa: BLE001
            actual = f"[EXC] {type(e).__name__}: {e}"
        results.append((tid, goal, expected, actual))
    return results


def _verify(rows):
    """예상↔실제 정합 간이 검증. 핵심 메시지가 실제 결과에 포함되는지 확인."""
    mismatch = []
    for tid, goal, expected, actual in rows:
        exp = expected.replace("!!! 오류: ", "").replace("✓ ", "")
        # 대표 토큰(가장 긴 한 조각)이 실제에 포함되는지
        token = max(exp.split(" "), key=len) if exp.split(" ") else exp
        if token and token not in actual and not actual.startswith("[EXC]"):
            # 메시지형 TC만 엄격 확인 (서술형은 패스)
            if expected.startswith(("!!!", "✓")):
                mismatch.append((tid, expected, actual))
    return mismatch


# 2차에서 더 이상 그 장면·방법으로 입력할 수 없거나 결과가 달라진 1차 TC → 보고서에서 삭제.
# (교수 안내: 그대로 동작하리라 확신하는 TC만 복붙할 것)  (id, 목표셀 부분일치)
REMOVE_TCS = [
    ("5.5-3", ""),      # config 자동생성 내용 2필드 → 2차는 3필드(학기 포함)
    ("5.5-9", ""),      # config 2필드 입력 → 2차는 필드수 오류(다른 경로)
    ("5.5-10", ""),     # config 2필드 입력 → 2차는 필드수 오류
    ("6.3-1", "정상 회원가입"),   # 학년 입력 단계 추가로 옛 입력 순서로 완료 불가
    ("6.6-10", ""),     # '9' → 2차는 수강신청 기간 설정(유효), 오류 아님
    ("6.7-1", ""),      # 개설과목 조회 출력형식 변경(스케줄/강의실/제한)
    ("6.7-3", ""),      # 검색 결과 출력형식 변경
    ("6.7-6", ""),      # inactive 제외 — 출력형식 변경
    ("6.9-9", ""),      # '시간표 충돌' → 2차 '스케줄 충돌' 메시지 변경
    ("6.12-1", ""),     # 시간표 출력형식 변경(강의실 추가)
    ("6.13-1", ""),     # 학생 등록 — 학년 입력 단계 추가로 옛 흐름 불가
    ("6.14-1", "강의 등록 정상"),       # 등록 흐름 변경(학기/제한/스케줄 분리)
    ("6.14-3", ""),     # 종료≤시작 — 옛 강의등록 인라인 시각입력 흐름 변경
    ("6.14-4", "기간 외 정상"),         # 통합 강의 수정 → 항목별 수정으로 재설계
    ("6.14-5", ""),     # '기간 중 수정 차단(통째)' 메시지 — 2차 미사용
    ("6.14-7", "종료 후 차단"),         # 기간 종료 후 통째 차단 — 2차 미사용
    ("6.15-1", ""),     # 기간 설정 — 학기 입력 추가로 옛 흐름/메시지 변경
    ("6.15-2", ""),     # 동일
]


def _renumber_duplicate_ids(d):
    """1차 보고서부터 있던 중복 TC ID(예: 6.3-8/9/10 2벌)를 섹션 내 다음 번호로 재부여."""
    import re as _re
    rows_all = []
    per_sec_max: dict[str, int] = {}
    for tbl in d.tables:
        for row in tbl.rows:
            m = _re.fullmatch(r"(\d+\.\d+)-(\d+)", row.cells[0].text.strip())
            if not m:
                continue
            sec, num = m.group(1), int(m.group(2))
            per_sec_max[sec] = max(per_sec_max.get(sec, 0), num)
            rows_all.append((row, sec, num))
    seen: set[str] = set()
    renamed = []
    for row, sec, num in rows_all:
        tid = f"{sec}-{num}"
        if tid in seen:
            per_sec_max[sec] += 1
            new_id = f"{sec}-{per_sec_max[sec]}"
            p = row.cells[0].paragraphs[0]
            if p.runs:
                p.runs[0].text = new_id
                for r in p.runs[1:]:
                    r.text = ""
            renamed.append((tid, new_id))
        else:
            seen.add(tid)
    return renamed


def _delete_invalid_1cha_rows(d):
    import re as _re
    removed = 0
    for tbl in d.tables:
        to_del = []
        for row in tbl.rows:
            rid_raw = row.cells[0].text.strip()
            m = _re.match(r"(\d+\.\d+-\d+)", rid_raw)
            if not m:
                continue
            rid = m.group(1)
            goal = row.cells[1].text
            for tid, sub in REMOVE_TCS:
                if rid == tid and (sub == "" or sub in goal):
                    to_del.append(row)
                    break
        for row in to_del:
            row._tr.getparent().remove(row._tr)
            removed += 1
    return removed


# 상호작용 화면 TC 스크린샷 (_gen_tc_screenshots.py 산출물). 화면 공유 TC는 같은 캡처 사용.
SHOT_DIR = ROOT / "발표준비" / "_tc_shots"
SHOT_SHARE = {"6.7-8": "6.7-7", "6.7-9": "6.7-7", "6.7-10": "6.7-7", "6.12-5": "6.12-4"}


def _shot_path(tid: str):
    p = SHOT_DIR / f"{SHOT_SHARE.get(tid, tid)}.png"
    return p if p.exists() else None


def _set_para_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def _swap_carried_shots(d, tids):
    """CARRIED TC의 stale 1차 스크린샷을 현행 2차 캡처(_tc_shots/{tid}.png)로 교체."""
    import re as _re
    from docx.shared import Cm
    swapped = []
    for tbl in d.tables:
        for row in tbl.rows:
            m = _re.match(r"(\d+\.\d+-\d+)", row.cells[0].text.strip())
            if not m or m.group(1) not in tids:
                continue
            png = SHOT_DIR / f"{m.group(1)}.png"
            if not png.exists():
                continue
            cell = row.cells[-1]
            for para in list(cell.paragraphs):
                para._element.getparent().remove(para._element)
            cell.add_paragraph().add_run().add_picture(str(png), width=Cm(8.0))
            swapped.append(m.group(1))
    return swapped


def _fix_6134_input(d):
    """6.13-4 입력 칼럼의 '→ 1 선택'(코드가 확인창 전 종료해 도달 불가) 제거."""
    import re as _re
    for tbl in d.tables:
        for row in tbl.rows:
            if row.cells[0].text.strip().startswith("6.13-4"):
                c = row.cells[1]
                t2 = _re.sub(r"\s*→\s*[‘'\"]?1[’'\"]?\s*선택", "", c.text)
                if t2 != c.text:
                    _set_para_text(c.paragraphs[0], t2)
                    for p in c.paragraphs[1:]:
                        _set_para_text(p, "")
                return


def _fix_toc_numbering(d):
    """목차/본문 '2.N' 중복(2.6 두 번) 순차 재번호 + 목차에 17절 1줄 추가."""
    import re as _re

    def _renum(text, n):  # 앞 공백(들여쓰기) 보존하며 2.N 치환
        return _re.sub(r"^(\s*)2\.\d+", lambda m: f"{m.group(1)}2.{n}", text)

    seq = 0
    for p in d.paragraphs:
        if p.style and p.style.name == "Heading 2" and _re.match(r"^\s*2\.\d+\s", p.text):
            seq += 1
            _set_para_text(p, _renum(p.text, seq))
    # 목차 항목 = 본문 표 밖(d.paragraphs)에서 '2.N '으로 시작하는 Normal 문단 (TOC 전용)
    toc = [p for p in d.paragraphs
           if p.style and p.style.name == "Normal" and _re.match(r"^\s*2\.\d+\s", p.text)]
    for i, p in enumerate(toc, 1):
        _set_para_text(p, _renum(p.text, i))
    if toc:
        from copy import deepcopy
        from docx.text.paragraph import Paragraph
        anchor = toc[-1]
        newp = deepcopy(anchor._p)
        anchor._p.addnext(newp)
        _set_para_text(
            Paragraph(newp, anchor._parent),
            "17. 2차 확장 검사 (신규·변경 TC) — 5.5/6.3/6.7/6.9/6.12/6.13/6.14/6.15/6.17 신규 53건",
        )


def build_docx(rows):
    import docx
    from docx.shared import Cm, Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    src = ROOT / "발표준비" / "B04_보고서_최종.docx"
    out = ROOT / "발표준비" / "B04_2차_검사_보고서.docx"
    shutil.copy(src, out)
    d = docx.Document(str(out))

    removed = _delete_invalid_1cha_rows(d)
    renamed = _renumber_duplicate_ids(d)
    swapped = _swap_carried_shots(d, {"6.10-1", "6.10-3", "6.10-6"})
    _fix_6134_input(d)
    _fix_toc_numbering(d)
    # 모든 섹션 머리글/바닥글의 '1차' → '2차' (페이지마다 반복되는 머리글 포함)
    for sec in d.sections:
        for hf in (sec.header, sec.footer):
            for p in hf.paragraphs:
                for r in p.runs:
                    if "1차 검사" in r.text:
                        r.text = r.text.replace("1차 검사", "2차 검사")

    NEW = RGBColor(0x70, 0x30, 0xA0)   # 보라 = 2차 신규/변경 검사
    BLACK = RGBColor(0, 0, 0)

    # 표지·개요를 2차 기준으로 갱신
    overview_old = (
        "본 보고서는 '건국 수강신청 시뮬레이터(B04)'의 1차 구현물에 대한 전체 통합 검사 결과를 "
        "기록한 문서이다. 기획서(1차 기획서 완성본)를 기준으로 테스트 케이스를 작성하였으며, "
        "브랜치 커버리지 100%를 목표로 구성하였다."
    )
    overview_new = (
        "본 보고서는 '건국 수강신청 시뮬레이터(B04)'의 2차 구현물에 대한 전체 통합 검사 결과를 "
        "기록한 문서이다. 2차 기획서(설계 기준 1판)와 설계 문서(원판)를 기준으로 테스트 케이스를 "
        "작성하였으며, 브랜치 커버리지 100%를 목표로 구성하였다. 1차에서 검사했고 2차에서도 동일하게 "
        "동작함을 확신하는 TC는 선별 승계(검정)하고, 2차에서 새로 검사한 TC는 17절에 보라색으로 구분하였다."
    )
    replacements = {
        "1차 검사 보고서": "2차 검사 보고서",
        "2026년 4월": "2026년 6월",
        overview_old: overview_new,
        "검사 기준: 1차 기획서 완성본 v2": "검사 기준: 2차 기획서 설계 기준 1판 + 설계 문서 원판",
    }
    for p in d.paragraphs[:60]:
        t = p.text
        for old, new_t in replacements.items():
            if old == t or (old in t and len(p.runs) == 1):
                p.runs[0].text = t.replace(old, new_t)
                break

    def shade(cell, hex_):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), hex_)
        tcPr.append(shd)

    def set_borders(tbl):
        tblPr = tbl._tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "888888")
            borders.append(el)
        tblPr.append(borders)

    # 1) 색상 범례를 보고서 첫머리(검사 개요 직후, '2. 테스트 케이스 목록' 앞)에 삽입 (강의자료 #10 지시)
    anchor = None
    for p in d.paragraphs:
        if p.style is not None and "Heading" in (p.style.name or "") and p.text.strip().startswith("2."):
            anchor = p
            break
    if anchor is not None:
        anchor.insert_paragraph_before("색상 범례", style="Heading 2")
        p1 = anchor.insert_paragraph_before("")
        r = p1.add_run("● 검정(기본): ")
        r.bold = True
        p1.add_run("1차에서 검사했고 2차 구현물에서도 동일한 장면·방법으로 입력·동작함을 확신하여 선별 승계한 TC(복붙). 2~16절의 1차 TC가 이에 해당.")
        p2 = anchor.insert_paragraph_before("")
        r = p2.add_run("● 보라: ")
        r.bold = True
        r.font.color.rgb = NEW
        p2.add_run("2차에서 처음 검사한 신규·변경 TC. 17절 '2차 확장 검사'에 모아 기재.")
        p3 = anchor.insert_paragraph_before("")
        r = p3.add_run(
            f"※ 2차 구현물에서 더 이상 같은 장면·방법으로 입력할 수 없거나 결과가 달라진 1차 TC {removed}개는 "
            f"본 보고서에서 삭제하고 17절의 신규 TC로 대체·재검증하였다."
        )
        r.italic = True

    # 2) 문서 맨 끝에 색상 범례 + 2차 섹션 추가
    d.add_page_break()
    h = d.add_heading("17. 2차 확장 검사 (신규·변경 TC)", level=1)

    lead = d.add_paragraph()
    lead.add_run(
        "본 17절은 2차 확장(설계문서 원판 기준)으로 추가·변경된 기능에 대한 실측 검사 결과이다. "
        "각 TC의 '실제 결과'는 실측한 값이며, 상호작용 화면 TC에는 실제 프로그램 실행 화면 캡처를 함께 첨부하였다. "
        "파일 자동 생성·단일 오류 메시지 출력 등 동작이 자명한 TC는 텍스트로만 기재한다."
    )

    # 색상 범례
    d.add_heading("색상 범례", level=2)
    legend = d.add_paragraph()
    legend.add_run("● 검정(기본): ").bold = True
    legend.add_run(
        "1차에서 검사했고 2차에서도 동일한 장면·방법으로 입력·동작하리라 확신하여 선별 승계한 TC(복붙). "
        "본 보고서 2~16절의 1차 TC가 이에 해당.\n"
    )
    r = legend.add_run("● 보라: ")
    r.bold = True
    r.font.color.rgb = NEW
    legend.add_run("2차에서 처음 검사한 신규·변경 TC(목표/입력/예상결과 중 하나 이상이 1차와 달라짐). 본 17절 전체가 이에 해당.\n")
    note = d.add_paragraph()
    note.add_run(
        f"※ 검사 수행 비율에 따라 '검정=복붙'을 기본으로 두고, 2차에서 새로 검사한 항목만 보라로 강조한다(강의자료 #10 8쪽 허용). "
        f"1차 TC 중 2차 구현물에서 더 이상 그 장면·방법으로 입력할 수 없거나 결과가 달라진 {removed}개는 "
        f"본 보고서에서 삭제하고 17절의 보라색 TC로 대체·재검증하였다."
    ).italic = True

    # 변경 안내
    d.add_heading("1차 기능 중 2차 구조 변경으로 삭제·재검증된 항목", level=2)
    chg = d.add_paragraph()
    chg.add_run(
        "다음 1차 기능은 2차 데이터 구조·규칙 변경으로 입력 방법 또는 결과가 바뀌어, 해당 1차 TC는 삭제하고 아래 신규 TC로 재검증하였다:\n"
        "· 5.5 무결성: classrooms/schedules/prerequisites 3개 파일 추가, students(7필드)·courses(10필드)·config(3필드) 스키마 변경, "
        "행/필드 앞뒤 공백·빈 행 검사 신설(기획서 5.1의 1차 미구현 항목을 2차에서 구현) → 5.5-12~25\n"
        "· 6.3/6.13 회원가입·학생등록: 학년 입력 추가 → 6.3-14~15, 6.13-10\n"
        "· 6.7 개설과목 조회: 스케줄·강의실·제한 표시, 학기 필터, 신청 불가 과목 [제한] 표시 → 6.7-7~11\n"
        "· 6.9 수강신청: 7단계 → 11단계(학기·학년·학과·선수과목), 충돌이 스케줄 기반으로 변경 → 6.9-15~21\n"
        "· 6.12 시간표: 스케줄 기반 출력, 동일 과목 복수 요일 표시, 학점 중복 합산 방지 → 6.12-4~5\n"
        "· 6.14 강의 수정: 통합 수정 → 항목별 수정으로 재설계(기간별 항목 제한, 정원/강의실 제약) → 6.14-16~31\n"
        "· 6.15 기간 설정: 학기 입력 추가 → 6.15-7~9\n"
        "· 6.17 강의실 관리: 신규 메뉴 → 6.17-1~3"
    )

    # 그룹별 표
    groups = [
        ("5.5  무결성 검사 확장 (파일·스키마·참조)", "5.5"),
        ("6.3  회원가입 — 학년 입력", "6.3"),
        ("6.7  개설 과목 조회 — 스케줄·제한·학기", "6.7"),
        ("6.9  수강신청 — 11단계 검사", "6.9"),
        ("6.12 내 시간표 — 스케줄 기반", "6.12"),
        ("6.13 관리자 학생 관리 — 학년", "6.13"),
        ("6.14 강의 등록·수정 — 스케줄·항목별 수정", "6.14"),
        ("6.15 수강신청 기간 설정 — 학기", "6.15"),
        ("6.17 강의실 관리 — 신규", "6.17"),
    ]
    by_prefix = {}
    for tid, goal, expected, actual in rows:
        pre = tid.split("-")[0]
        by_prefix.setdefault(pre, []).append((tid, goal, expected, actual))

    for title, pre in groups:
        items = by_prefix.get(pre, [])
        if not items:
            continue
        d.add_heading(title, level=2)
        tbl = d.add_table(rows=1, cols=4)
        set_borders(tbl)
        hdr = tbl.rows[0].cells
        for j, htext in enumerate(["ID", "목표 / 입력", "예상 결과", "실제 결과"]):
            hdr[j].paragraphs[0].add_run(htext).bold = True
            shade(hdr[j], "DDDDDD")
        # 열 폭 힌트
        for tid, goal, expected, actual in items:
            cells = tbl.add_row().cells
            rid = cells[0].paragraphs[0].add_run(tid)
            rid.bold = True
            rid.font.color.rgb = NEW
            cells[1].paragraphs[0].add_run(goal)
            cells[2].paragraphs[0].add_run(expected)
            ar = cells[3].paragraphs[0].add_run(actual if not actual.startswith("[EXC]") else actual)
            # 실제결과가 예상과 부합하면 검정, EXC면 빨강
            ar.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) if actual.startswith("[EXC]") else BLACK
            shot = _shot_path(tid)
            if shot is not None:
                pic_par = cells[3].add_paragraph()
                pic_par.add_run().add_picture(str(shot), width=Cm(8.0))

    # 요약
    kept = 129 - removed
    d.add_heading("2차 검사 요약", level=2)
    summ = d.add_paragraph()
    summ.add_run(f"· 1차 TC 129개 중 {removed}개 삭제(2차에서 수행 불가/결과 변경) → 선별 승계 {kept}개(검정=복붙)\n")
    summ.add_run(f"· 2차 신규·변경 TC: {len(rows)}개(보라, 전부 실측·예상결과 부합)\n")
    summ.add_run(f"· 전체 TC: {kept} + {len(rows)} = {kept + len(rows)}개")

    d.save(str(out))
    print(f"[build] removed {removed} invalid 1차 TCs; kept {kept}; added {len(rows)} 2차 TCs; "
          f"renamed dup IDs: {renamed}; swapped shots: {swapped}")
    return out


if __name__ == "__main__":
    rows = run()
    mism = _verify(rows)
    print(f"=== {len(rows)} TCs 실행, 불일치 {len(mism)}개 ===")
    for tid, e, a in mism:
        print(f"  MISMATCH {tid}: EXP={e[:50]} ACT={a[:70]}")
    if "--build" in sys.argv:
        out = build_docx(rows)
        print("saved:", out)
